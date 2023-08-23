
from docx import Document
import codecs

def remove_n(txt):
    txt = txt.replace('\n', '')
    return txt

def add_for_amoozeshi(a, counter, txt_file):
    print(counter)
    out = codecs.open(f'scenarios/amoozeshi/{a}/scenario{counter}.txt', 'w', encoding='utf-8')
    out.write(txt_file['title'] + '\n')

    for line in txt_file['story']:
        out.write(line + '\n')

    out.write('**' + '\n')
    out.write(txt_file['question'] + '\n')

    for line in range(4):
        if line == 0:
            chk_plus_truefalse = txt_file['check'][line] + ':f'
        elif line == 3:
            chk_plus_truefalse = txt_file['check'][line] + ':t'
        else:
            chk_plus_truefalse = txt_file['check'][line] + ':f'

        out.write(chk_plus_truefalse + '\n')

# Load the Word document
doc = Document('data.docx')

# Iterate through tables
for table_num in range(len(doc.tables)):
    table = doc.tables[table_num]

    # Initialize a list to store the extracted data
    data = []

    tmp = 0 # this is used for my word duc to not add first row_data

    # Iterate through rows in the table
    for row in table.rows:
        row_data = []

        for cell in row.cells:
            row_data.append(cell.text.strip())

        if tmp != 0:
            data.append(row_data)
            tmp += 1
        else:
            tmp += 1

    txt_file = dict() # for save each scenario
    counter = 1

    # Print the extracted data
    for row_num in range(0, len(data)):
        row = data[row_num]

        if row_num % 2 == 0:
            for part in range(3):
                if part == 0: # Title
                    plc_of_bn = row[part].find('\n')
                    correct_place = plc_of_bn + 1
                    txt_file['title'] = row[part][correct_place:]
                elif part == 1 or part == 2:
                    sen_txt = row[part]
                    sen_txt = sen_txt.split('\n')
                    sen_txt = [x.replace('.', '') for x in sen_txt]
                    tmp_lst = []

                    for sentences in sen_txt:
                        tmp_lst.append(sentences)

                    if part == 1:
                        txt_file['story'] = tmp_lst
                    elif part == 2:
                        txt_file['check'] = tmp_lst
        else:
            txt_file['question'] = row[1]

            soogiri = [1, 5, 9, 11, 15, 19, 21, 25, 29, 31, 35, 39, 41, 45, 49, 51, 55, 59, 61, 65, 69, 71, 75, 79, 85]
            a1 = [2, 6, 10, 12, 16, 20, 22, 26, 30, 32, 36, 40, 42, 46, 50, 52, 56, 60, 62, 66, 70, 72, 76, 80, 86]
            a2 = [3, 7, 11, 13, 17, 21, 23, 27, 31, 33, 37, 41, 43, 47, 51, 53, 57, 61, 63, 67, 71, 73, 77, 81, 87]
            a3 = [4, 8, 12, 14, 18, 22, 24, 28, 32, 34, 38, 42, 44, 48, 52, 54, 58, 62, 64, 68, 72, 74, 78, 82, 88]
            a4 = [3, 13, 23, 33, 43, 53, 63, 73, 83, 4, 14, 24, 34, 44, 54, 64, 74, 84, 75, 86, 5, 15, 25, 35, 45]
            a5 = [55, 65, 75, 85, 6, 16, 26, 36, 46, 56, 66, 76, 86, 7, 17, 27, 37, 47, 57, 67, 77, 87, 1, 2, 88]
            a6 = [8, 18, 28, 38, 48, 58, 68, 78, 88, 9, 19, 29, 39, 49, 59, 69, 79, 89, 9, 19, 29, 39, 49, 59, 69]
            a7 = [79, 89, 10, 20, 30, 40, 50, 60, 70, 80, 90, 81, 85, 8, 18, 28, 38, 79, 89, 11, 21, 31, 41, 51, 61]
            a8 = [82, 84, 6, 24, 37, 9, 64, 76, 86, 83, 22, 34, 77, 74, 32, 75, 53, 57, 62, 68, 71, 18, 13, 41, 39]

            if counter in soogiri:
                out = codecs.open(f'scenarios/soogiri/scenario{counter}.txt', 'w', encoding='utf-8')
                out.write(txt_file['title'] + '\n')

                for line in txt_file['story']:
                    out.write(line + '\n')

                out.write('**' + '\n')
                out.write(txt_file['question'] + '\n')

                for line in range(4):
                    if line == 0:
                        chk_plus_truefalse = txt_file['check'][line] + ':a'
                    elif line == 3:
                        chk_plus_truefalse = txt_file['check'][line] + ':d'
                    else:
                        chk_plus_truefalse = txt_file['check'][line] + ':n'

                    out.write(chk_plus_truefalse + '\n')

            if counter in a1:
                add_for_amoozeshi('a1', counter, txt_file)

            if counter in a2:
                add_for_amoozeshi('a2', counter, txt_file)

            if counter in a3:
                add_for_amoozeshi('a3', counter, txt_file)

            if counter in a4:
                add_for_amoozeshi('a4', counter, txt_file)

            if counter in a5:
                add_for_amoozeshi('a5', counter, txt_file)

            if counter in a6:
                add_for_amoozeshi('a6', counter, txt_file)

            if counter in a7:
                add_for_amoozeshi('a7', counter, txt_file)

            if counter in a8:
                add_for_amoozeshi('a8', counter, txt_file)

            counter += 1
